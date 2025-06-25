import tkinter as tk
import os
from tkinter import messagebox
from tkinter import ttk
import datetime
from tkcalendar import DateEntry
from docx import Document 

def display_data(name, reason, date_ocurred, location_ocurred):
    print(f"Name: {name}")
    print(f"Reason: {reason}")
    print(f"Date Occurred: {date_ocurred}")
    print(f"Location: {location_ocurred}")

def formatted_name(name):
    value = name.strip().split()

    if len(value) < 2:
        return name
    last_name_index = -1

    if value[-1].lower() == "jr" or value[-1].lower() == "jr.":
        last_name_index=-2
    last_name = value[last_name_index]
    rest = " ".join(value[:last_name_index])
    formatted_name = f"{last_name}, {rest}"
    return formatted_name


def create_folder(name):
    folder_name = formatted_name(name)

    base_path = os.path.join("\\\\ENTERIPADDRESSHERE\\Shared\\Clients", "Potential Clients")
    full_path = os.path.join(base_path, folder_name)

    os.makedirs(base_path, exist_ok=True)
    os.makedirs(full_path, exist_ok=True)

    print(f"Folder created at: {full_path}")

    return full_path

def create_file(name, reason, date, location, folder_path):
    doc = Document()
    doc.add_heading(f'Client Info: {name}', level=1)

    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Reason for Calling: {reason}")
    doc.add_paragraph(f"Date Ocurred: {date}")
    doc.add_paragraph(f"Location Ocurred: {location}")
    # doc.add_paragraph(f"{}")
    # doc.add_paragraph(f"{}")
    doc.add_paragraph(f"Date Contacted {datetime.datetime.now().strftime('%m-%d-%Y %H:%M')}")

    file_path = os.path.join(folder_path, f"{name.replace(' ', '_')}_intake.docx")
    doc.save(file_path)

    print (f"Saved Word document at {file_path}")

def generate_UI():
    window = tk.Tk()
    custom_reason_entry = None
    window.title("Client Intake Form")    
    window_width = int(window.winfo_screenwidth() * .6)
    window_height = int(window.winfo_screenheight() * .6)
    window.geometry(f"{window_width}x{window_height}")

    name_label = tk.Label(window, text="Name:")
    name_label.grid(row=0, column=0)
    name_entry = tk.Entry(window)
    name_entry.grid(row=0, column=1)

    reason_label = tk.Label(window, text="Reason:")
    reason_label.grid(row=1, column=0)
    reasons = [
        "Assault",
        "Burglary",
        "Driving under the influence (DUI)",
        "Driving while intoxicated (DWI)",
        "Driving while suspended (DWS)",
        "Domestic Abuse",
        "Drug Possession",
        "Parole",
        "Probation",
        "Sex Crimes",
        "Traffic Violation",
        "Other"
    ]
    reasons_variable = tk.StringVar()
    reasons_variable.set("Select a reason")

    #reason_entry = tk.Entry(window) #used to be the way to type in an entry, the more control the better
    #reason_entry = tk.OptionMenu(window, reasons_variable, *reasons) #used as a dropdown menu
    reason_entry = ttk.Combobox(window, textvariable=reasons_variable, values=reasons, state = "readonly")
    reason_entry.grid(row=1, column=1)

    def reason_handler(event):
        nonlocal custom_reason_entry

        if reasons_variable.get().strip() =="Other":
            if not custom_reason_entry:
                custom_reason_entry = tk.Entry(window)
                custom_reason_entry.grid(row=1, column=2, padx=5)
                custom_reason_entry.insert(0, "Enter Custom Reason")
            else:
                if custom_reason_entry:
                    custom_reason_entry.destroy()
                    custom_reason_entry = None

    reason_entry.bind("<<ComboboxSelected>>", reason_handler)


    date_occurred_label = tk.Label(window, text="Date Occurred:")
    date_occurred_label.grid(row=2, column=0)
    date_occurred_entry = DateEntry(window, date_pattern='mm-dd-yyyy')
    #date_occurred_entry = tk.Entry(window) #ability to enter date, this way it is formatted the same eveyr time
    date_occurred_entry.grid(row=2, column=1)

    location_label = tk.Label(window, text="Location:")
    location_label.grid(row=3, column=0)
    location_entry = tk.Entry(window)
    location_entry.grid(row=3, column=1)

    def submit():
        nonlocal custom_reason_entry
        name = name_entry.get().strip()
        reason = reasons_variable.get().strip()
        date = date_occurred_entry.get().strip()
        location = location_entry.get().strip()

        if not name:
            name = "NO NAME"
        if reason == "Other" and custom_reason_entry:
            reason = custom_reason_entry.get().strip()
            if not reason:
                reason = "Not specified"
        if not date:
            date = "NO DATE"
        if not location:
            location = "NO LOCATION"

        display_data(name, reason, date, location)
        if name:
            create_file(name, reason, date, location, create_folder(name))

        name_entry.delete(0, tk.END)
        reasons_variable.set("Select a reason")
        date_occurred_entry.delete(0, tk.END)
        location_entry.delete(0, tk.END)

        if custom_reason_entry:
            custom_reason_entry.destroy()
            custom_reason_entry = None

        messagebox.showinfo("Success!", "Data submitted successfully")
    
    submit_button = tk.Button(window, text="Submit", command=submit)
    submit_button.grid(row=4, column=0, columnspan=2)

    window.mainloop()



if __name__ == "__main__":
    generate_UI()
    
